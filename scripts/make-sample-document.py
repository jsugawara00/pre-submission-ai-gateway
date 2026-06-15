# -*- coding: utf-8 -*-
"""
動作確認用の貿易書類サンプルを sample-document/ に生成する。
1件の貿易取引として、荷送人・荷受人・商品・数量・金額・B/L番号などの
照合キーを全書類で統一する（整合した正常セット。不一致は仕込まない）。

- インボイス      → Excel (.xlsx)
- パッキングリスト → Excel (.xlsx)
- 船荷証券        → Word  (.docx)
- 原産地証明書    → Word  (.docx)

すべて架空。実在企業名・実在B/L番号は使用しない。
"""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "sample-document")

# ===== 1件分の取引データ（照合キーの正本。全書類でこれを参照する）=====
D = {
    "exporter": "Saigon Textile Export Co., Ltd.",
    "exporter_addr": "123 Nguyen Hue Blvd, District 1, Ho Chi Minh City, Vietnam",
    "importer": "Tokyo Apparel Import K.K.",
    "importer_addr": "4-5-6 Nihonbashi, Chuo-ku, Tokyo 103-0027, Japan",
    "invoice_no": "INV-2026-0418",
    "invoice_date": "2026-04-18",
    "po_no": "PO-TA-26-0331",
    "lc_no": "N/A (T/T)",
    "vessel": "OCEAN HARMONY V.025E",
    "bl_no": "SGNTYO260418",
    "bl_date": "2026-04-20",
    "pol": "Ho Chi Minh City, Vietnam",
    "pod": "Tokyo, Japan",
    "final_dest": "Tokyo, Japan",
    "origin": "Vietnam",
    "incoterms": "CIF Tokyo",
    "payment": "T/T 30 days after B/L date",
    "currency": "USD",
    "description": "Men's Cotton T-Shirt",
    "style_no": "CT-100",
    "hs_code": "6109.10",
    "qty": 5000,            # pcs
    "unit_price": 4.50,     # USD/pc
    "amount": 22500.00,     # USD
    "cartons": 100,
    "pcs_per_carton": 50,
    "net_weight": 1000.0,   # kg
    "gross_weight": 1150.0, # kg
    "measurement": 8.5,     # CBM
}

thin = Side(style="thin", color="000000")
BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
HEAD_FILL = PatternFill("solid", fgColor="D9E1F2")


def _title(ws, text):
    ws.merge_cells("A1:F1")
    c = ws["A1"]
    c.value = text
    c.font = Font(size=16, bold=True)
    c.alignment = Alignment(horizontal="center")


def make_invoice():
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"
    for col, w in zip("ABCDEF", (16, 22, 14, 12, 12, 14)):
        ws.column_dimensions[col].width = w
    _title(ws, "COMMERCIAL INVOICE")

    rows = [
        ("Exporter:", D["exporter"], "", "Invoice No:", D["invoice_no"], ""),
        ("", D["exporter_addr"], "", "Invoice Date:", D["invoice_date"], ""),
        ("Consignee:", D["importer"], "", "P/O No:", D["po_no"], ""),
        ("", D["importer_addr"], "", "Payment:", D["payment"], ""),
        ("Vessel:", D["vessel"], "", "Incoterms:", D["incoterms"], ""),
        ("Port of Loading:", D["pol"], "", "B/L No:", D["bl_no"], ""),
        ("Port of Discharge:", D["pod"], "", "Country of Origin:", D["origin"], ""),
    ]
    r = 3
    for row in rows:
        for i, val in enumerate(row):
            c = ws.cell(row=r, column=i + 1, value=val)
            if i in (0, 3):
                c.font = Font(bold=True)
        r += 1

    r += 1
    headers = ["HS Code", "Description / Style", "Quantity (pcs)", "Unit Price", "Amount", "Currency"]
    for i, h in enumerate(headers):
        c = ws.cell(row=r, column=i + 1, value=h)
        c.font = Font(bold=True)
        c.fill = HEAD_FILL
        c.border = BORDER
        c.alignment = Alignment(horizontal="center")
    r += 1
    line = [
        D["hs_code"],
        f'{D["description"]} (Style {D["style_no"]})',
        D["qty"],
        D["unit_price"],
        D["amount"],
        D["currency"],
    ]
    for i, val in enumerate(line):
        c = ws.cell(row=r, column=i + 1, value=val)
        c.border = BORDER
        if i == 2:
            c.number_format = "#,##0"
        if i in (3, 4):
            c.number_format = "#,##0.00"
    r += 1
    ws.cell(row=r, column=4, value="TOTAL:").font = Font(bold=True)
    tc = ws.cell(row=r, column=5, value=D["amount"])
    tc.font = Font(bold=True)
    tc.number_format = "#,##0.00"
    ws.cell(row=r, column=6, value=D["currency"]).font = Font(bold=True)

    r += 3
    ws.cell(row=r, column=1, value=f'Total packages: {D["cartons"]} cartons')
    ws.cell(row=r + 1, column=1, value=f'Net weight: {D["net_weight"]:,.1f} kg / Gross weight: {D["gross_weight"]:,.1f} kg')
    ws.cell(row=r + 3, column=5, value=D["exporter"]).font = Font(bold=True)
    ws.cell(row=r + 4, column=5, value="Authorized Signature")

    wb.save(os.path.join(OUT_DIR, "01_Commercial_Invoice.xlsx"))


def make_packing_list():
    wb = Workbook()
    ws = wb.active
    ws.title = "Packing List"
    for col, w in zip("ABCDEFG", (10, 24, 12, 14, 14, 14, 12)):
        ws.column_dimensions[col].width = w
    _title(ws, "PACKING LIST")

    rows = [
        ("Exporter:", D["exporter"], "", "Invoice No:", D["invoice_no"]),
        ("Consignee:", D["importer"], "", "Invoice Date:", D["invoice_date"]),
        ("Vessel:", D["vessel"], "", "B/L No:", D["bl_no"]),
        ("Port of Loading:", D["pol"], "", "Port of Discharge:", D["pod"]),
    ]
    r = 3
    for row in rows:
        for i, val in enumerate(row):
            c = ws.cell(row=r, column=i + 1, value=val)
            if i in (0, 3):
                c.font = Font(bold=True)
        r += 1

    r += 1
    headers = ["Carton No", "Description / Style", "Qty/Ctn", "Cartons", "Total Pcs", "N.W. (kg)", "G.W. (kg)"]
    for i, h in enumerate(headers):
        c = ws.cell(row=r, column=i + 1, value=h)
        c.font = Font(bold=True)
        c.fill = HEAD_FILL
        c.border = BORDER
        c.alignment = Alignment(horizontal="center")
    r += 1
    line = [
        f'1-{D["cartons"]}',
        f'{D["description"]} (Style {D["style_no"]})',
        D["pcs_per_carton"],
        D["cartons"],
        D["qty"],
        D["net_weight"],
        D["gross_weight"],
    ]
    for i, val in enumerate(line):
        c = ws.cell(row=r, column=i + 1, value=val)
        c.border = BORDER
        if i in (2, 3, 4):
            c.number_format = "#,##0"
        if i in (5, 6):
            c.number_format = "#,##0.1"
    r += 1
    ws.cell(row=r, column=3, value="TOTAL:").font = Font(bold=True)
    for col, val, fmt in ((4, D["cartons"], "#,##0"), (5, D["qty"], "#,##0"),
                          (6, D["net_weight"], "#,##0.1"), (7, D["gross_weight"], "#,##0.1")):
        c = ws.cell(row=r, column=col, value=val)
        c.font = Font(bold=True)
        c.number_format = fmt
        c.border = BORDER

    r += 2
    ws.cell(row=r, column=1, value=f'Total measurement: {D["measurement"]} CBM').font = Font(bold=True)
    ws.cell(row=r + 2, column=5, value=D["exporter"]).font = Font(bold=True)
    ws.cell(row=r + 3, column=5, value="Authorized Signature")

    wb.save(os.path.join(OUT_DIR, "02_Packing_List.xlsx"))


def _kv(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(str(value))
    return p


def make_bill_of_lading():
    doc = Document()
    h = doc.add_heading("BILL OF LADING", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("(Non-negotiable copy — for sample / testing use only)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(9)

    _kv(doc, "B/L No", D["bl_no"])
    _kv(doc, "B/L Date", D["bl_date"])
    _kv(doc, "Shipper", f'{D["exporter"]}, {D["exporter_addr"]}')
    _kv(doc, "Consignee", f'{D["importer"]}, {D["importer_addr"]}')
    _kv(doc, "Notify Party", D["importer"])
    _kv(doc, "Ocean Vessel / Voyage", D["vessel"])
    _kv(doc, "Port of Loading", D["pol"])
    _kv(doc, "Port of Discharge", D["pod"])
    _kv(doc, "Place of Delivery", D["final_dest"])

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Marks & Numbers", "No. of Pkgs", "Description of Goods", "Gross Weight / Meas."]):
        cell.paragraphs[0].add_run(text).bold = True
    row = table.add_row().cells
    row[0].text = f'{D["importer"]}\n{D["invoice_no"]}\nMADE IN VIETNAM'
    row[1].text = f'{D["cartons"]} CARTONS'
    row[2].text = f'{D["description"]} (Style {D["style_no"]})\n{D["qty"]:,} PCS\nHS Code: {D["hs_code"]}'
    row[3].text = f'{D["gross_weight"]:,.1f} KG\n{D["measurement"]} CBM'

    doc.add_paragraph()
    _kv(doc, "Freight", D["incoterms"] + " (Freight Prepaid)")
    _kv(doc, "Number of Original B/L", "THREE (3)")
    doc.add_paragraph()
    doc.add_paragraph("SHIPPED on board the vessel named above in apparent good order and condition.")
    sign = doc.add_paragraph("\n\nFor the Carrier")
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(os.path.join(OUT_DIR, "03_Bill_of_Lading.docx"))


def make_certificate_of_origin():
    doc = Document()
    h = doc.add_heading("CERTIFICATE OF ORIGIN", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("(Sample / testing use only)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(9)

    _kv(doc, "Certificate No", f'CO-{D["invoice_no"]}')
    _kv(doc, "Exporter", f'{D["exporter"]}, {D["exporter_addr"]}')
    _kv(doc, "Consignee", f'{D["importer"]}, {D["importer_addr"]}')
    _kv(doc, "Country of Origin", D["origin"])
    _kv(doc, "Means of Transport", f'By sea — {D["vessel"]}')
    _kv(doc, "Port of Loading", D["pol"])
    _kv(doc, "Port of Discharge", D["pod"])
    _kv(doc, "Invoice No / Date", f'{D["invoice_no"]} / {D["invoice_date"]}')

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells,
                          ["Item", "Marks & Nos", "Description of Goods", "HS Code", "Quantity"]):
        cell.paragraphs[0].add_run(text).bold = True
    row = table.add_row().cells
    row[0].text = "1"
    row[1].text = f'{D["cartons"]} CARTONS'
    row[2].text = f'{D["description"]} (Style {D["style_no"]})'
    row[3].text = D["hs_code"]
    row[4].text = f'{D["qty"]:,} PCS'

    doc.add_paragraph()
    doc.add_paragraph(
        "It is hereby certified that the goods described above are of "
        f'{D["origin"]} origin.'
    )
    doc.add_paragraph()
    doc.add_paragraph("Place and Date: Ho Chi Minh City, " + D["invoice_date"])
    sign = doc.add_paragraph("\nAuthorized Signature / Stamp")
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(os.path.join(OUT_DIR, "04_Certificate_of_Origin.docx"))


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    make_invoice()
    make_packing_list()
    make_bill_of_lading()
    make_certificate_of_origin()
    print("生成完了:")
    for f in sorted(os.listdir(OUT_DIR)):
        print("  -", f)


if __name__ == "__main__":
    main()
